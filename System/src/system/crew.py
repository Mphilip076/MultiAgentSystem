from crewai import Agent, Crew, Process, Task, LLM
from crewai.project import CrewBase, agent, crew, task
from crewai.agents.agent_builder.base_agent import BaseAgent
from pydantic import BaseModel, Field
from crewai_tools import SerperDevTool
import os
import base64
from email.message import EmailMessage
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from crewai.tools import tool
import requests
import io
from docx import Document
import docx
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Initialize the tools
search_tool = SerperDevTool()


# Set the scope to only sending emails
SCOPES = ['https://www.googleapis.com/auth/gmail.send']

CUR_DIR = os.path.dirname(os.path.abspath(__file__))
CREDENTIALS_PATH = os.path.join(CUR_DIR, '..', '..', 'credentials.json')
TOKEN_PATH = os.path.join(CUR_DIR, '..', '..', 'token.json')

def get_gmail_service():
    """Handles OAuth2 authentication and returns the Gmail service."""
    creds = None
    # token.json stores the user's access and refresh tokens
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)
    
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        # This will open a browser window for the first run
        if not os.path.exists(CREDENTIALS_PATH):
            raise FileNotFoundError(f"Credentials file not found at {os.path.abspath(CREDENTIALS_PATH)}. Please ensure it exists in the System/ directory.")
            
        flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_PATH, SCOPES)
        creds = flow.run_local_server(port=0)
        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())
    
    return build('gmail', 'v1', credentials=creds)

@tool("Email Sender Tool")
def send_email_tool(subject: str, email_body: str, report_content: str) -> str:
    """Use this tool to send the finalized strategic report to the executive team 
    via the Gmail API. 
    Args:
        subject: The exact subject line for the email.
        email_body: The summarized email content (title, company, date, summary, and 3 bullet points).
        report_content: The full formatted text content of the report to attach as a Word document.
    """
    recipient_email = "mateoviteri13579@gmail.com"
    
    try:
        service = get_gmail_service()
        message = EmailMessage()
        
        # Structure the email
        message.set_content(email_body)
        message['To'] = recipient_email
        message['From'] = "me"
        message['Subject'] = subject

        def insert_hyperlink(paragraph, text, url):
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000EE')
            rPr.append(color)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            new_run.append(rPr)
            
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            new_run.append(text_elem)
            
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        def parse_and_add_links(paragraph, text):
            pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
            last_end = 0
            for match in pattern.finditer(text):
                start, end = match.span()
                if start > last_end:
                    paragraph.add_run(text[last_end:start])
                link_text = match.group(1)
                url = match.group(2)
                insert_hyperlink(paragraph, link_text, url)
                last_end = end
                
            if last_end < len(text):
                paragraph.add_run(text[last_end:])

        # Create Word Document attachment
        doc = Document()
        for paragraph in report_content.split('\n'):
            paragraph = paragraph.strip().replace('**', '')
            if not paragraph:
                continue
                
            if paragraph.startswith('### '):
                p = doc.add_heading(level=3)
                parse_and_add_links(p, paragraph[4:])
            elif paragraph.startswith('## '):
                p = doc.add_heading(level=2)
                parse_and_add_links(p, paragraph[3:])
            elif paragraph.startswith('# '):
                p = doc.add_heading(level=1)
                parse_and_add_links(p, paragraph[2:])
            elif paragraph.startswith('- ') or paragraph.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                parse_and_add_links(p, paragraph[2:])
            else:
                p = doc.add_paragraph()
                parse_and_add_links(p, paragraph)
                
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Attach the document
        message.add_attachment(
            doc_io.read(), 
            maintype='application', 
            subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', 
            filename='report.docx'
        )

        # Gmail API requires the message to be base64url encoded
        encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
        create_message = {'raw': encoded_message}
        
        # Execute the send command
        send_message = (service.users().messages().send(userId="me", body=create_message).execute())
        
        return f"Email successfully sent to {recipient_email}! Message ID: {send_message['id']}"
    
    except Exception as e:
        return f"Failed to send email. Error: {str(e)}"

@tool("URL Checker Tool")
def url_check_tool(url: str) -> str:
    """Use this tool to check if a URL is valid and doesn't return a 404 error.
    Args:
        url: The full URL string to check.
    """
    try:
        response = requests.head(url, timeout=5, allow_redirects=True, headers={'User-Agent': 'Mozilla/5.0'})
        if response.status_code >= 400:
            # Fallback to GET if HEAD fails
            response = requests.get(url, timeout=5, stream=True, headers={'User-Agent': 'Mozilla/5.0'})
            response.close()
            
        if response.status_code >= 400:
            return f"Invalid URL (Status {response.status_code}): {url}"
        return f"Valid URL: {url}"
    except Exception as e:
        return f"Failed to connect to URL. Error: {str(e)}"

class ReportTemplate(BaseModel):
    what_happened: str = Field(description="A clear and concise summary of the core event or news.")
    competitive_impact: str = Field(description="Detailed analysis of how this impacts the competitive landscape.")
    why_it_matters: str = Field(description="Explanation of the strategic significance and alignment with company goals.")
    tell_me_more: str = Field(description="The actual verified facts and detailed background context extracted from the validated data.")
    outlook: str = Field(description="Actionable strategic recommendations, future outlook, and competitor activities to monitor.")

@CrewBase
class System():
    """System crew"""

    agents: list[BaseAgent]
    tasks: list[Task]

    # --- AGENTS --- #
    @agent
    def researcher(self) -> Agent:
        return Agent(
            config=self.agents_config['researcher'], 
            verbose=True,
            tools=[search_tool],
            llm=LLM(model=os.getenv("MODEL"), base_url=os.getenv("OLLAMA_API_BASE"))
        )

    @agent
    def data_validator(self) -> Agent:
        return Agent(
            config=self.agents_config['data_validator'], 
            verbose=True,
            tools=[url_check_tool],
            llm=LLM(model=os.getenv("MODEL"), base_url=os.getenv("OLLAMA_API_BASE"))
        )

    @agent
    def report_creator(self) -> Agent:
        return Agent(
            config=self.agents_config['report_creator'], 
            verbose=True,
            llm=LLM(model=os.getenv("MODEL"), base_url=os.getenv("OLLAMA_API_BASE"))
        )

    @agent
    def report_validator(self) -> Agent:
        return Agent(
            config=self.agents_config['report_validator'], 
            verbose=True,
            llm=LLM(model=os.getenv("MODEL"), base_url=os.getenv("OLLAMA_API_BASE"))
        )

    @agent
    def send_report(self) -> Agent:
        return Agent(
            config=self.agents_config['send_report'], 
            verbose=True,
            tools=[send_email_tool],
            llm=LLM(model=os.getenv("MODEL"), base_url=os.getenv("OLLAMA_API_BASE"))
        )

    # --- TASKS --- #
    @task
    def research_task(self) -> Task:
        return Task(config=self.tasks_config['research_task'])

    @task
    def data_validation_task(self) -> Task:
        return Task(config=self.tasks_config['data_validation_task'])

    @task
    def report_creation_task(self) -> Task:
        return Task(
            config=self.tasks_config['report_creation_task'],
            output_pydantic=ReportTemplate 
        )

    @task
    def report_validation_task(self) -> Task:
        return Task(config=self.tasks_config['report_validation_task'])

    @task
    def send_report_task(self) -> Task:
        return Task(config=self.tasks_config['send_report_task'])

    @crew
    def crew(self) -> Crew:
        """Creates the System crew"""
        return Crew(
            agents=self.agents, 
            tasks=self.tasks, 
            verbose=True,
            process=Process.sequential
        )
