from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field
import os
import base64
from email.message import EmailMessage
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
import io
from docx import Document
import docx
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# Set the scope to only sending emails
SCOPES = ['https://www.googleapis.com/auth/gmail.send']

CUR_DIR = os.path.dirname(os.path.abspath(__file__))
CREDENTIALS_PATH = os.path.join(CUR_DIR, '..', '..', '..', 'credentials.json')
TOKEN_PATH = os.path.join(CUR_DIR, '..', '..', '..', 'token.json')

def get_gmail_service():
    """Handles OAuth2 authentication and returns the Gmail service."""
    creds = None
    # token.json stores the user's access and refresh tokens
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)
    
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        # This will open a browser window for the first run
        if not os.path.exists(CREDENTIALS_PATH):
            raise FileNotFoundError(f"Credentials file not found at {os.path.abspath(CREDENTIALS_PATH)}. Please ensure it exists in the System/ directory.")
            
        flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_PATH, SCOPES)
        creds = flow.run_local_server(port=0)
        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())
    
    return build('gmail', 'v1', credentials=creds)

class EmailSenderToolInput(BaseModel):
    """Input schema for EmailSenderTool."""
    subject: str = Field(..., description="The exact subject line for the email.")
    email_body: str = Field(..., description="The summarized email content (title, company, date, summary, and 3 bullet points).")

class EmailSenderTool(BaseTool):
    name: str = "Email Sender Tool"
    description: str = (
        "Use this tool to send the finalized strategic report to the executive team via the Gmail API."
    )
    args_schema: Type[BaseModel] = EmailSenderToolInput

    def _run(self, subject: str, email_body: str) -> str:
        recipient_email = "mateoviteri13579@gmail.com"
        
        try:
            with open('final_report.md', 'r', encoding='utf-8') as f:
                report_content = f.read()
                
            service = get_gmail_service()
            message = EmailMessage()
            
            # Structure the email
            message.set_content(email_body)
            message['To'] = recipient_email
            message['From'] = "me"
            message['Subject'] = subject

            def insert_hyperlink(paragraph, text, url):
                part = paragraph.part
                r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)

                new_run = OxmlElement('w:r')
                
                rPr = OxmlElement('w:rPr')
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '0000EE')
                rPr.append(color)
                
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                
                new_run.append(rPr)
                
                text_elem = OxmlElement('w:t')
                text_elem.text = text
                new_run.append(text_elem)
                
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)

            def parse_and_add_links(paragraph, text):
                pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
                last_end = 0
                for match in pattern.finditer(text):
                    start, end = match.span()
                    if start > last_end:
                        paragraph.add_run(text[last_end:start])
                    link_text = match.group(1)
                    url = match.group(2)
                    insert_hyperlink(paragraph, link_text, url)
                    last_end = end
                    
                if last_end < len(text):
                    paragraph.add_run(text[last_end:])

            # Create Word Document attachment
            doc = Document()
            for paragraph in report_content.split('\n'):
                paragraph = paragraph.strip().replace('**', '')
                if not paragraph:
                    continue
                    
                if paragraph.startswith('### '):
                    p = doc.add_heading(level=3)
                    parse_and_add_links(p, paragraph[4:])
                elif paragraph.startswith('## '):
                    p = doc.add_heading(level=2)
                    parse_and_add_links(p, paragraph[3:])
                elif paragraph.startswith('# '):
                    p = doc.add_heading(level=1)
                    parse_and_add_links(p, paragraph[2:])
                elif paragraph.startswith('- ') or paragraph.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    parse_and_add_links(p, paragraph[2:])
                else:
                    p = doc.add_paragraph()
                    parse_and_add_links(p, paragraph)
                    
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            # Attach the document
            message.add_attachment(
                doc_io.read(), 
                maintype='application', 
                subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', 
                filename=f'{datetime.now().strftime("%d%m%y")}-Alert-{subject}.docx'
            )

            # Gmail API requires the message to be base64url encoded
            encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
            create_message = {'raw': encoded_message}
            
            # Execute the send command
            send_message = (service.users().messages().send(userId="me", body=create_message).execute())
            
            # Clean up the intermediate file
            try:
                os.remove('final_report.md')
            except OSError:
                print("Warning: Failed to remove final_report.md")
                
            return f"Email successfully sent to {recipient_email}! Message ID: {send_message['id']}"
        
        except Exception as e:
            return f"Failed to send email. Error: {str(e)}"
